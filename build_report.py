from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "期末项目报告-完成版.docx"


def set_font(run, name="宋体", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def clear_doc(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=10.5, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        set_cell_shading(cell, fill)


def add_para(doc, text="", style=None, first_line=True):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_font(r, size=11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, name="黑体", size=15 if level == 1 else 12.5, bold=True, color=(31, 78, 121))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r, size=11)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, width in enumerate(widths):
        table.columns[idx].width = Cm(width)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, fill="E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    for style_name in ["List Bullet", "List Number"]:
        s = doc.styles[style_name]
        s.font.name = "宋体"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        s.font.size = Pt(11)


def build():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    r = title.add_run("天气数据可视化分析平台")
    set_font(r, name="黑体", size=18, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    r = sub.add_run("《数据可视化》课程期末项目报告")
    set_font(r, name="宋体", size=12, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(12)
    r = meta.add_run("姓名：待填写    学号：待填写    小组：2人")
    set_font(r, size=11)

    add_heading(doc, "一、项目简介")
    add_para(
        doc,
        "本项目以 2025 年 5 月至 2026 年 5 月的城市天气历史数据为基础，完成了一个网页形式的天气数据可视化分析平台。平台围绕“平均情况总览、各省详情、全球视野”三个入口组织信息：总览页用于快速比较不同城市的平均最高温、平均最低温、空气质量和天气类型；各省详情页用于查看中国不同省份及城市的气温、天气、风向等局部特征；全球视野页则通过三维地球和城市散点展示国内外城市在同一时间维度下的天气状态。",
    )
    add_para(
        doc,
        "项目的主要目标不是简单罗列天气表格，而是把时间、空间和天气类型三个维度结合起来，让使用者能够从宏观平均、区域细节和全球分布三个层次理解数据。页面整体采用深色界面、蓝色交互高亮和统一的导航按钮，使不同页面之间保持视觉一致性，同时提高在课堂展示时的辨识度。",
    )

    add_heading(doc, "二、数据介绍")
    add_para(
        doc,
        "数据主要包括中国 34 个城市以及部分境外城市的历史天气记录。国内数据按月份保存为 weather_data2505.csv 至 weather_data2605.csv，并额外补充香港、澳门、台北三个城市的独立文件；国际数据汇总在 weather_202505_202605.csv 中。原始字段包括日期、城市、城市代码、最高温、最低温、天气状况、风向风力、空气质量指数等，其中境外数据不统一提供空气质量字段，因此在后续处理与展示中对字段差异进行了兼容。",
    )
    add_bullets(
        doc,
        [
            "时间范围：2025-05 至 2026-05，能够覆盖春夏秋冬四季变化。",
            "空间范围：中国主要省会及代表城市，另含香港、澳门、台北和若干国际城市。",
            "核心指标：最高温、最低温、天气状况、风向风力、空气质量指数、城市经纬度。",
            "派生数据：city_average_data.json 保存城市平均指标和天气类型计数，供总览页快速加载。",
        ],
    )

    add_heading(doc, "三、数据处理")
    add_para(
        doc,
        "数据处理主要分为采集、清洗、统一字段、聚合统计和页面适配五个步骤。采集阶段使用 Python 脚本按城市代码和月份批量请求天气数据，并将不同月份保存为 CSV 文件；清洗阶段统一中英文字段名，将“最高温、最低温、天气”等字段重命名为 max_temperature、min_temperature、weather_condition 等页面更容易读取的字段；聚合阶段统计每个城市的平均最高温、平均最低温、平均 AQI、晴/多云/雨/雪等天气类型出现次数，并输出 JSON 供页面直接调用。",
    )
    add_numbered(
        doc,
        [
            "城市代码匹配：通过 cities.csv 和 match_city_code.py 建立城市名称与天气接口城市代码的对应关系。",
            "月度数据爬取：gainData_csv.py 和相关脚本按月份生成国内城市天气 CSV，gainData_csv_inter.py 负责境外城市数据采集。",
            "字段清洗与兼容：处理普通城市、港澳台城市和境外城市字段差异，缺失 AQI 的数据在展示时采用默认或不展示策略。",
            "统计聚合：data-processor.html 对所有 CSV 进行读取，按城市累计温度、AQI、天气类型和风向出现次数，形成 city_average_data.json。",
            "可视化适配：为三维地球补充城市经纬度字典，并把天气文本映射为图标、颜色和提示框内容。",
        ],
    )

    add_heading(doc, "四、可视化设计")
    add_para(
        doc,
        "本项目采用 HTML、CSS、JavaScript、D3、ECharts 和 Three.js 进行开发。整体结构以 index_final.html 作为主要整合页面，global_vision.html 作为全球视野页面，并通过统一的顶部导航完成页面切换。设计上尽量减少说明性文字，把交互重点放在筛选控件、地图/图表区域和悬浮提示信息上。",
    )
    add_heading(doc, "1. 平均情况总览", level=2)
    add_para(
        doc,
        "平均情况总览页面面向快速比较，重点展示各城市在一年周期内的平均最高温、平均最低温、空气质量和天气类型占比。页面使用柱状图、折线图、饼图等图表形态，让用户可以快速识别温度较高、温差明显或空气质量较差的城市。该页承担整个项目的“入口仪表盘”功能，因此强调信息密度和整体可读性。",
    )
    add_heading(doc, "2. 各省详情", level=2)
    add_para(
        doc,
        "各省详情页突出空间选择与局部分析，使用中国地图或省份/城市列表帮助用户进入具体区域。用户选择城市后，可以查看对应城市的气温变化、天气类型构成、风向统计和空气质量等信息。该页的设计重点在于把全国层面的平均比较下钻到城市层面，支撑更细粒度的天气特征分析。",
    )
    add_heading(doc, "3. 全球视野", level=2)
    add_para(
        doc,
        "全球视野页最初考虑使用 ECharts GL 的 globe 和 scatter3D 实现 3D 地球散点，但在贴图、交互控制和悬浮效果调试过程中，最终改为 Three.js 实现。页面使用真实地球贴图和高度贴图，配合 OrbitControls 实现类似建模软件的自由旋转、缩放和视角控制；城市天气点采用金色微尖头标记，避免普通球形散点显得杂乱；悬浮时只加深当前点颜色，并显示城市、天气图标、天气文字和温度范围。",
    )
    add_para(
        doc,
        "日期筛选从单一下拉框优化为“年、月、日”三级选择，提高了查找具体日期的效率。数据读取方面，全球视野页同时加载国内月度 CSV、港澳台 CSV 和国际 CSV，并在统一的 mapRow 函数中进行字段兼容，再根据当前日期刷新三维地球上的标记。",
    )

    add_heading(doc, "五、分工说明")
    add_table(
        doc,
        ["成员", "主要工作", "具体说明"],
        [
            [
                "成员1",
                "数据采集与处理；各省详情页面；全球视野页面优化",
                "负责爬取并整理天气数据，处理城市代码匹配、字段清洗和 CSV/JSON 输出；编写各省详情页面的数据读取与城市下钻逻辑；后期参与全球视野页的交互、散点样式和页面效果优化。",
            ],
            [
                "成员2",
                "平均情况总览页；全球视野页；页面整合与优化",
                "负责总览页图表结构和平均指标展示；搭建全球视野页的三维地球基础结构；统一三个页面的导航、深色主题、按钮样式和展示逻辑，并完成最终页面整合。",
            ],
        ],
        [2.0, 4.2, 9.0],
    )

    add_heading(doc, "六、大模型辅助使用说明")
    add_para(
        doc,
        "项目开发过程中使用大模型主要进行方案咨询、架构选择和交互细节优化，而不是直接照搬生成代码。整理 AI 交互记录后，报告中保留偏“解决方案/架构建议”的部分，排除了大段完整代码内容。大模型辅助主要体现在以下方面：",
    )
    add_bullets(
        doc,
        [
            "咨询 3D 地球实现方案，对比 ECharts GL 与 Three.js 在贴图、旋转控制、散点交互方面的适用性。",
            "讨论全球视野页的视觉风格，包括写实地球贴图、均匀光照、金色城市标记和自定义 Tooltip。",
            "优化日期筛选方式，将单一下拉框调整为年、月、日三级联动，提高筛选效率。",
            "梳理多来源 CSV 的字段兼容方法，使国内、港澳台和国际城市数据可以在同一页面中读取。",
            "根据课堂展示需求调整页面导航和深色 UI 风格，使三个功能页保持一致。",
        ],
    )

    add_heading(doc, "七、结果分析")
    add_para(
        doc,
        "从可视化结果看，中国不同城市的气温差异具有明显地域特征：南方及沿海城市平均最低温较高，北方和高海拔地区季节温差更明显；天气类型方面，多云和降雨在多数城市中占比较高，晴天数量在北方部分城市更突出；空气质量指数在不同城市之间存在差异，适合通过总览页进行横向比较。全球视野页则把国内外城市放到同一三维空间中，直观呈现同一天不同地区的天气差异，增强了项目展示的空间感和互动性。",
    )
    add_para(
        doc,
        "项目最终形成了从“数据采集-清洗聚合-多页面展示-交互分析”的完整流程。平均情况总览适合快速发现整体规律，各省详情适合查看局部差异，全球视野适合进行空间展示和课堂汇报。三类页面互相补充，使项目既有数据分析价值，也有较好的展示效果。",
    )

    add_heading(doc, "八、课程小结")
    add_para(
        doc,
        "通过本次项目，我对数据可视化的理解从“画出图表”进一步扩展到“围绕问题组织数据和交互”。在开发过程中，最重要的收获是：可视化并不是图表越多越好，而是要根据数据维度选择合适的表达方式。平均指标适合用柱状图或折线图比较，城市分布适合地图或三维场景，天气类型则适合用颜色、图标和提示框降低理解成本。",
    )
    add_para(
        doc,
        "项目中也遇到了一些问题，例如不同来源数据字段不统一、部分城市缺少经纬度、三维地球贴图和光照效果调试成本较高、页面整合后样式容易不一致等。通过逐步清洗数据、封装字段映射、统一页面导航和反复调试交互，最终基本完成了预期功能。后续如果继续完善，可以加入更多城市的经纬度、增加时间轴动画、提供按天气类型筛选的功能，并进一步优化移动端适配。",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
